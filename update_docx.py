#!/usr/bin/env python3
"""
更新论文 docx：
1. 效果部分：加入 CenterPoint++ Epoch 30 各指标表格 + mAP 计算方式说明
2. 方法部分：加入混合训练泛化性先验 + 小样本度量学习叙述
3. 插入最优推理效果图
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DOCX_PATH = 'work/基于无锚框 3D 目标检测的仓储机器人实时语义感知1.docx'
IMG_PATH = 'output/best_inference_viz/002981.png'
OUTPUT_PATH = 'work/基于无锚框 3D 目标检测的仓储机器人实时语义感知1.docx'  # overwrite

doc = Document(DOCX_PATH)


# ========== 辅助函数 ==========
def set_cell_shading(cell, color="D9E2F3"):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_paragraph_after(paragraph, text, style='Normal', bold=False, font_size=None):
    """在给定段落后插入新段落"""
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    # Actually, python-docx doesn't support this easily.
    # We'll use the paragraph's parent body instead
    pass


def insert_paragraph_after_element(element, text, bold=False, font_size=None, color=None, style='Normal'):
    """在给定元素后插入段落（使用 lxml）"""
    from lxml import etree
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    new_p = OxmlElement('w:p')
    # 添加 run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)
    element.addnext(new_p)
    return new_p


def get_table_text(table):
    """打印表格内容（调试用）"""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            print(f"  [{i},{j}]: {cell.text[:50]}")


# ========== 1. 更新「效果」部分：添加指标表格和 mAP 说明 ==========
# 找到 "效果" heading
effect_heading = None
for p in doc.paragraphs:
    if p.style.name.startswith('Heading') and '效果' in p.text:
        effect_heading = p
        break

if effect_heading is None:
    raise ValueError("未找到「效果」标题")

# 找到效果部分末尾（下一个 Heading）
effect_end = None
started = False
for p in doc.paragraphs:
    if p == effect_heading:
        started = True
        continue
    if started and p.style.name.startswith('Heading'):
        effect_end = p
        break

# 在效果标题后，现有内容之前插入 mAP 表格
# 先找到效果标题后的第一个段落
effect_paragraphs = []
started = False
for p in doc.paragraphs:
    if p == effect_heading:
        started = True
        continue
    if started and p.style.name.startswith('Heading'):
        break
    if started:
        effect_paragraphs.append(p)

# 清理效果部分旧内容（除了数据集描述等保留）
# Strategy: 直接在效果标题后插入新内容，保留原有段落不变，在它们之前插入

# 构建 mAP 表格数据
map_headers = ['类别', '3D AP@0.50 (mAP)', 'BEV AP@0.50', 'BBox AP@0.50', 'AOS']
map_data = [
    ['箱子', '81.82%', '85.89%', '86.16%', '85.88%'],
    ['电动运输车', '66.89%', '84.77%', '84.77%', '84.69%'],
    ['货运自行车',  '2.27%',  '2.27%',  '2.27%',  '2.27%'],
    ['无人搬运车', '25.49%', '72.22%', '85.05%', '85.02%'],
    ['叉车',       '0.00%',  '0.00%', '17.30%',  '9.88%'],
    ['**均值**',   '**35.30%**', '**49.03%**', '**55.11%**', '**53.55%**'],
]

# 在效果标题后插入提示文字和表格
p_intro = insert_paragraph_after_element(
    effect_heading._element,
    '表X 展示了 CenterPoint++ 模型在最优 epoch（第 30 轮）上的各类别检测精度指标。'
    '所有 AP 值基于 KITTI 标准评估协议，采用 11 点插值法计算，IoU 判别阈值为 3D 交并比 ≥ 0.50。',
    bold=False, font_size=10.5
)

# Insert a table
table = doc.add_table(rows=len(map_data) + 1, cols=len(map_headers), style='Table Grid')
# 但 doc.add_table 会把表格追加到文档末尾...需要移动

# 更好的方法：用 python-docx 在特定位置插入表格比较复杂
# 改用：直接操作 XML

# Let me rebuild the docx approach: clear paragraphs between effect_heading and next heading,
# then insert new content directly

# === Simplified approach: just replace the content in existing effect paragraphs ===
# P39 = 效果 heading
# P40 onward = effect content paragraphs

# Find paragraph indices
para_idx_map = {}
for i, p in enumerate(doc.paragraphs):
    para_idx_map[p] = i

# P40: existing effect intro text
# We'll replace and add content here

# Clear existing effect paragraphs' text (keep structure)
# For paragraphs between effect heading and next heading
effect_range_start = para_idx_map.get(effect_heading, -1) + 1
effect_range_end = para_idx_map.get(effect_end, len(doc.paragraphs))

# 修改 P40 (效果第一段): 替换为 mAP 说明
doc.paragraphs[effect_range_start].text = (
    '为验证本文提出基于 Anchor-free 3D 目标检测系统的有效性，我们在仓库仿真数据集中进行了全面测试。'
    '评估采用 KITTI 官方评估协议，使用 11 点插值平均精度（11-point Interpolation Average Precision）'
    '作为核心指标。'
)

# 在 P40 后插入 mAP 计算说明
p40_elem = doc.paragraphs[effect_range_start]._element
next_elem = p40_elem.getnext()

# Helper to insert paragraph at specific position
def insert_para(text, after_elem, bold_start=None, font_size=10.5):
    from lxml import etree
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)
    after_elem.addnext(new_p)
    return new_p

# Insert mAP computation methodology
current_pos = p40_elem
current_pos = insert_para(
    'mAP 计算方式：评估系统遵循 KITTI 3D 目标检测标准协议。对于每一类别，首先通过匈牙利匹配算法将预测框与'
    '真值框进行关联，匹配条件为 3D 旋转交并比（Rotated 3D IoU）≥ 阈值（对自定义类别统一设为 0.50）。'
